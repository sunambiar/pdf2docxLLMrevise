
import pdfplumber
from docx import Document
from language_tool_python import LanguageTool
import re

# Initialize the grammar correction tool with English rules
grammar_tool = LanguageTool('en-US')

#import pdfplumber
#from docx import Document
#from language_tool_python import LanguageTool
#import re

# Initialize the grammar correction tool with English rules
#grammar_tool = LanguageTool('en')

def extract_paragraphs_with_sanskrit(pdf_path):
    """Extract text paragraphs from a PDF, identifying Sanskrit verses"""
    with pdfplumber.open(pdf_path) as pdf:
        paragraphs = []

        for page in pdf.pages:
            extracted_text = page.extract_text()

            # Assuming Sanskrit text is detected by a specific pattern (you might need to refine this)
            sanskrit_pattern = r'\bDevanagari\b|\bSanskrit\b|(\u0900-\u097F|\uA015|\uA4DF)\s*'
            matches = re.findall(sanskrit_pattern, extracted_text)

            if matches:
                # If Sanskrit text is found, save the whole page as a single paragraph
                paragraphs.append(extracted_text)
            else:
                # Split by two newlines to separate paragraphs
                lines = extracted_text.split('\n\n')

                for line in lines:
                    if re.search(sanskrit_pattern, line):  # If line contains Sanskrit
                        continue

                    # Otherwise, split the paragraph at typical sentence breaks (e.g., periods)
                    sentences = re.split('[.!?]', line)
                    paragraphs.extend([s + '\n' for s in sentences if s])

    return paragraphs

def correct_non_sanskrit_paragraphs(paragraphs):
    """Correct grammar in non-Sanskrit verse paragraphs"""
    corrected = []
    for p in paragraphs:
        try:
            # Basic text cleaning before correction
            cleaned = " ".join(p.split())
            corrected_paragraph = grammar_tool.correct(cleaned)
            corrected.append(corrected_paragraph)
        except Exception as e:
            print(f"Error correcting non-Sanskrit paragraph: {e}")
            corrected.append(p)
    return corrected

#def save_to_word(doc, paragraphs, word_path):
def save_to_word(doc, paragraphs):    
    """Save the paragraphs to a Word document"""
    #doc = Document()
    for paragraph in paragraphs: 
        p = doc.add_paragraph(paragraph)
        # Optionally add spacing between paragraphs (if desired)
        if p and doc.paragraphs:
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.add_run().add_break()

    #doc.save(word_path)



# Example usage
#pdf_file = "your_document.pdf"
#word_file = "corrected_output.docx"

# 1. Extract paragraphs from PDF, identifying Sanskrit verses
#paragraphs = extract_paragraphs_with_sanskrit(pdf_file)
#print(f"Extracted {len(paragraphs)} total paragraphs (including non-Sanskrit)")

# 2. Correct grammar in non-Sanskrit paragraphs
#corrected_paragraphs = correct_non_sanskrit_paragraphs(paragraphs)
#print(f"Corrected {len(corrected_paragraphs)} non-Sanskrit paragraphs")

# 3. Save the result to a Word document (including preserved Sanskrit verses)
#save_to_word(corrected_paragraphs, word_file)
#print(f"Saved corrected document: {word_file}")


def vivekachudamani_correct(pdf_files, word_file):
    wordDoc = Document()
    for pdf_file in pdf_files:
       print(f"ReadFile: {pdf_file}")
       paragraphs = extract_paragraphs_with_sanskrit(pdf_file)
       corrected_paragraphs = correct_non_sanskrit_paragraphs(paragraphs)
       save_to_word(wordDoc, corrected_paragraphs)

    wordDoc.save(word_file)
    print(f"Output stored in {word_file}")


if __name__ == "__main__":
    genPath = "/Users/sureshnambiar/snPDFdocxRewrite/"

    wordPath = genPath + "output/"
    word_file = wordPath + "Vivekachudamani_corrected.docx"
    wordDoc = Document()

    pdfPath = genPath + "/Vivekachudamani/"

    for i in range(1, 5):                
       pdf_file = pdfPath + f"Vivekachudamani_Chinmayananda_part_{i}.pdf"
       print(f"ReadFile: {pdf_file}")

       paragraphs = extract_paragraphs_with_sanskrit(pdf_file)
       corrected_paragraphs = correct_non_sanskrit_paragraphs(paragraphs)

       save_to_word(wordDoc, corrected_paragraphs)

    wordDoc.save(word_file)
    print(f"Output stored in {word_file}")


